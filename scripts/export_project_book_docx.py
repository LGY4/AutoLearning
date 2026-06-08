from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("docs/generated/AutoLearning_V2_软件杯项目书.md")
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("docs/generated/AutoLearning_V2_软件杯项目书.docx")

IMAGE_RE = re.compile(r"^!\[(?P<alt>.*?)\]\((?P<path>[^)]+)\)$")
ORDERED_RE = re.compile(r"^\d+\.\s+")


def set_default_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)


def add_cover(document: Document) -> None:
    title_text = SOURCE.stem
    subtitle_text = "基于多智能体协同与讯飞生态的个性化智能学习平台"
    logo_path = SOURCE.parent / "assets" / "project_logo.png"

    if logo_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_path), width=Inches(1.8))

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title_text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(subtitle_text)
    sub_run.font.name = "Microsoft YaHei"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sub_run.font.size = Pt(13)

    for line in [
        "",
        "参赛队伍：待填写",
        "所在学校：待填写",
        "团队成员：待填写",
        "指导教师：待填写",
        "完成日期：2026年6月",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(12)

    document.add_section(WD_SECTION.NEW_PAGE)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return all(c and set(c) <= {"-", ":"} for c in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]] | None, int]:
    if start + 1 >= len(lines):
        return None, start

    header = lines[start].strip()
    separator = lines[start + 1].strip()
    if not (header.startswith("|") and header.endswith("|") and is_table_separator(separator)):
        return None, start

    rows: list[list[str]] = []
    index = start
    while index < len(lines):
        current = lines[index].strip()
        if not (current.startswith("|") and current.endswith("|")):
            break
        rows.append([cell.strip() for cell in current.strip("|").split("|")])
        index += 1

    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return

    header = rows[0]
    body = [row for row in rows[2:] if len(row) == len(header)]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text

    for row in body:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    document.add_paragraph("")


def add_image(document: Document, source_dir: Path, stripped: str) -> bool:
    match = IMAGE_RE.match(stripped)
    if not match:
        return False

    alt = match.group("alt").strip()
    raw_path = match.group("path").strip()
    image_path = (source_dir / raw_path).resolve()
    if not image_path.exists():
        document.add_paragraph(f"[缺失图片] {raw_path}")
        return True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))

    if alt:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(alt)
        cap_run.font.name = "Microsoft YaHei"
        cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        cap_run.font.size = Pt(10)
    return True


def add_code_block(document: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
    document.add_paragraph("")


def add_markdown(document: Document, markdown_text: str, source_dir: Path) -> None:
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped = raw_line.strip()

        if not stripped:
            document.add_paragraph("")
            index += 1
            continue

        if stripped == "---":
            document.add_paragraph("")
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        table_rows, next_index = parse_table(lines, index)
        if table_rows:
            add_table(document, table_rows)
            index = next_index
            continue

        if add_image(document, source_dir, stripped):
            index += 1
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            index += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            index += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped[2:].strip())
            index += 1
            continue

        if ORDERED_RE.match(stripped):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(ORDERED_RE.sub("", stripped, count=1))
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped[2:].strip())
            run.italic = True
            index += 1
            continue

        document.add_paragraph(stripped)
        index += 1


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE}")

    markdown = SOURCE.read_text(encoding="utf-8")

    document = Document()
    set_default_style(document)
    add_cover(document)
    add_markdown(document, markdown, SOURCE.parent)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(str(OUTPUT.resolve()))


if __name__ == "__main__":
    main()
