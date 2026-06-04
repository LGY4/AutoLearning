"""生成 AutoLearning 项目图解说明 docx 文件"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUT = os.path.dirname(os.path.abspath(__file__))
doc = Document()

# ── 样式设置 ─────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Microsoft YaHei'
    h.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

# ── 封面 ─────────────────────────────────────────────────────────────
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AutoLearning 项目\n智能学习系统图解说明')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('系统架构 · 时序交互 · 活动流程 · 协作泳道 · 状态机')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('生成日期: 2026-05-29\n技术栈: FastAPI + React + LangGraph + PostgreSQL + ChromaDB').font.size = Pt(10)

doc.add_page_break()

# ── 目录 ─────────────────────────────────────────────────────────────
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目智能性评估',
    '  1.1 代码实现角度的智能体现',
    '  1.2 实际使用角度的智能体现',
    '  1.3 智能化总结',
    '二、系统架构流程图',
    '三、自适应学习时序图',
    '四、资源生成活动图',
    '五、多智能体协作泳道图',
    '六、学习路径节点状态图',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 第一章：智能性评估 ────────────────────────────────────────────────
doc.add_heading('一、项目智能性评估', level=1)

doc.add_heading('1.1 代码实现角度的智能体现', level=2)

intel_items = [
    ('四维度画像模型',
     '每个知识点独立评估4个维度(mastery掌握/application应用/memory记忆/understanding理解)，'
     '每个维度三级(low/mid/high)，共81种组合。不同组合映射到不同教学策略，而非简单的"会/不会"二分法。'
     'composite_score 加权公式: 0.3*掌握 + 0.25*应用 + 0.2*记忆 + 0.25*理解。'),
    ('置信度感知合并',
     '画像更新采用置信度加权策略: 高置信度(>=0.7)直接覆盖，中置信度(0.4-0.7)加权平均，'
     '低置信度(<0.4)保留旧数据。避免低质量数据(如猜测的答题)污染高质量画像。'),
    ('规则+LLM混合架构',
     '策略引擎(strategy_engine)为纯规则引擎，81种四维度组合直接映射教学参数，无需LLM调用，响应<1ms。'
     '内容生成(document/quiz/mindmap等)使用LLM，兼顾速度与灵活性。'),
    ('自适应测验闭环',
     '辅导流程中的"先测后答"机制: 答对则难度递进(1->2->3)，答错则降级；'
     '四维度轮换测试(mastery->understanding->application->memory)；'
     '每次答题都触发画像更新闭环。'),
    ('多智能体LangGraph编排',
     '19个节点的StateGraph执行，支持条件路由、并行扇出(fan-out)、质量重试(质量分<0.7重试最多2次)、'
     'WebSearch兜底(生成失败时自动搜索网络资源)。非简单串行，是真正的有向无环图执行。'),
    ('知识图谱+图算法',
     'DAG结构的知识图谱，BFS回溯收集传递依赖 + 拓扑排序确保学习顺序正确。'
     '学"图"需要先学"树"，学"树"需要先学"链表"——自动解析依赖链。'),
    ('事件驱动画像更新',
     '10种事件类型覆盖所有画像变更来源: 诊断测验、自适应测验、练习评分、路径完成、错题记录、'
     '用户编辑、LLM提取、资源消费、对话行为、复习完成。每种事件有专用处理器。'),
    ('图感知RAG',
     '向量检索不仅看文本相似度，还扩展知识图谱上下文: 同节点的兄弟chunk也加入结果(分数打7折)。'
     '画像增强: 薄弱知识点在检索中加分+0.3，确保弱项相关内容优先。'),
]

for title, desc in intel_items:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}: ')
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.add_run(desc)

doc.add_heading('1.2 实际使用角度的智能体现', level=2)

usage_items = [
    ('冷启动诊断',
     '新用户首次使用时，系统通过3-5道定位题快速评估四维度，匹配到8个预定义学习者原型之一'
     '(如"视觉型初学者""理论强实践弱型")，用原型值初始化画像。无需等待大量数据积累。'),
    ('个性化资源推荐',
     '根据四维度+学习风格动态决定资源类型: 视觉型加mindmap/flowchart/video，'
     '动手型加code_case，阅读型加reading。不是所有人都看同样的内容。'),
    ('学习路径自动规划',
     '用户设定学习目标后，系统自动: 分析知识图谱 -> LLM规划路径 -> 图算法修正依赖顺序 -> '
     '创建节点状态机。完成一个节点自动解锁下一个。'),
    ('后台智能资源生成',
     '每次学习后(post_learning_update)，系统自动: 识别薄弱知识点 -> 规划缺失资源类型 -> '
     '后台线程生成(最多3个/周期)。用户无感知，下次登录时已有新资源可用。'),
    ('生成失败自动降级',
     '当LLM生成流程图/视频等资源失败时，自动触发DuckDuckGo/Bilibili网络搜索兜底，'
     '向用户展示可点击的网络资源链接，而非简单的"生成失败"。'),
    ('难度自适应调整',
     '答对题目 -> 难度+1，答错 -> 难度-1。维度轮换测试确保全面评估。'
     '连续答对3题后自动进入下一维度，避免单一维度过度测试。'),
    ('对话信号隐式提取',
     'LLM分析每轮对话，自动提取学习信号: confusion(困惑标记为弱项)、'
     'mastery_confirmed(确认理解微提升+0.1)、new_question(新延伸问题)。'
     '不需要用户显式做题也能更新画像。'),
    ('推荐系统自动刷新',
     '画像更新后自动清除推荐缓存并重新生成。推荐中包含"虚拟推荐"(suggested_generation)——'
     '建议生成但尚未存在的资源，可被后台自动消费生成。'),
]

for title, desc in usage_items:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}: ')
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    p.add_run(desc)

doc.add_heading('1.3 智能化总结', level=2)

summary_text = (
    'AutoLearning 的智能不是单点的 AI 调用，而是一个完整的多层反馈闭环系统:\n\n'
    '用户行为(对话/答题/资源消费) → 信号提取(对话信号/答题评分/行为分析) → '
    '画像更新(事件驱动/置信度加权/四维度评估) → 策略引擎(81种组合→教学参数) → '
    '执行层(资源生成/路径规划/辅导回答) → 质量控制(质量检查/重试/兜底) → '
    '用户行为(新一轮学习)\n\n'
    '核心创新: 四维度画像(非二分法) + 置信度感知合并 + 规则+LLM混合 + 图算法保障 + 多智能体协作'
)
doc.add_paragraph(summary_text)

doc.add_page_break()

# ── 第二章：系统架构流程图 ────────────────────────────────────────────
doc.add_heading('二、系统架构流程图', level=1)
doc.add_paragraph(
    '本图展示了 AutoLearning 系统的整体架构，从用户层到数据层的完整技术栈。'
)
doc.add_picture(os.path.join(OUT, '1_系统架构流程图.png'), width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('各层说明', level=2)
layers = [
    ('用户层 (蓝色)',
     '用户通过浏览器与系统交互，主要入口: ChatPanel(对话面板)、LearningMapPage(学习地图)、'
     'ResourceRenderer(资源渲染器)。'),
    ('前端层 (绿色)',
     'React + TypeScript 前端组件。ChatPanel 处理对话和测验交互；LearningMapPage 展示知识图谱和学习路径；'
     'ResourceRenderer 根据资源类型(document/quiz/mindmap/flowchart/video/code)分发到对应的渲染组件。'),
    ('API层 (橙色)',
     'FastAPI 后端接口。三个核心入口: /tutor/chat(辅导)、/learning/chat-stream(学习流SSE)、'
     '/resources/generate(资源生成)。所有接口通过依赖注入获取当前用户。'),
    ('智能体层 (紫色)',
     'MasterAgent 做意图识别与路由(6种意图: tutoring/resource_generation/learning_path/assessment/exercise/general_chat)；'
     'TutorService 处理辅导对话+自适应测验；LangGraph 编排19个节点的工作流。'),
    ('策略层 (紫色)',
     'ProfileService 管理四维度画像；StrategyEngine 81种组合映射教学参数；'
     'ResourcePlanner 综合弱项+目标+路径规划资源；AdaptiveService 执行学习后的自适应更新(8步操作)。'),
    ('资源生成层 (粉色)',
     '8种资源类型的独立生成器: Document(两阶段:大纲+填充)、Quiz(题库缓存+LLM生成)、'
     'Mindmap(Markmap格式)、Flowchart(draw.io XML)、Video(分镜脚本)、Code(完整可运行代码)。'),
    ('基础/数据层 (灰色)',
     'KnowledgeGraph(DAG知识图谱，15节点23边)、RAG/Embedding(ChromaDB向量检索+画像增强)、'
     'WebSearch(DuckDuckGo+Bilibili兜底搜索)。PostgreSQL持久化+ChromaDB向量库+LLM Provider。'),
]

for name, desc in layers:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{name}: ')
    run_t.font.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── 第三章：自适应学习时序图 ──────────────────────────────────────────
doc.add_heading('三、自适应学习时序图', level=1)
doc.add_paragraph(
    '本图展示了用户一次完整学习交互的时序流程，从发送问题到画像更新和资源推荐的全过程。'
    '分为三个阶段: 辅导阶段、测验评分阶段、自适应更新阶段。'
)
doc.add_picture(os.path.join(OUT, '2_自适应学习时序图.png'), width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('关键交互说明', level=2)
seq_items = [
    ('辅导阶段',
     '用户发送问题 → ChatPanel调用TutorService → TutorService识别意图(tutoring) → '
     '调用StrategyEngine获取教学参数(根据四维度) → 加载用户画像 → 返回辅导回答+附带测验题。'
     '教学参数直接影响回答的难度、风格和深度。'),
    ('测验评分阶段',
     '用户回答测验 → TutorService调用grading_service进行LLM语义评分(非选择题) → '
     'evaluate_knowledge_point更新四维度 → emit_event(ADAPTIVE_QUIZ)触发画像事件。'),
    ('自适应更新阶段',
     'post_learning_update执行8步操作: 画像更新→策略计算→资源推荐→缓存失效→'
     '战略资源规划→自动路径生成→目标资源生成→消费建议推荐。'
     'ResourcePlanner返回规划列表如[quiz, mindmap, video]，后台线程异步生成。'),
]
for title, desc in seq_items:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}: ')
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.add_run(desc)

doc.add_page_break()

# ── 第四章：资源生成活动图 ────────────────────────────────────────────
doc.add_heading('四、资源生成活动图', level=1)
doc.add_paragraph(
    '本图展示了资源生成的完整活动流程，从触发到质量控制到最终保存。'
    '核心特征: 并行扇出(Fan-out)——8种资源类型同时生成，每种独立走LLM生成→质量检查→保存的流程。'
)
doc.add_picture(os.path.join(OUT, '3_资源生成活动图.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('活动流程说明', level=2)
act_items = [
    ('触发入口 (3种)',
     '① 用户显式请求: "帮我生成一个快速排序的流程图"；'
     '② 画像驱动: post_learning_update中strategic_auto_generate根据弱项自动规划；'
     '③ 目标驱动: 设定学习目标后自动生成相关资源。'),
    ('意图识别与画像加载',
     'MasterAgent识别意图(resource_generation) → ProfileService加载画像 → '
     'StrategyEngine根据四维度+学习风格计算推荐参数 → ResourcePlanner综合弱项/目标/路径规划资源列表。'),
    ('并行扇出',
     '资源规划器确定要生成的类型后，LangGraph同时启动多个生成节点(如gen_document + gen_quiz + gen_mindmap)。'
     '每个节点独立执行，互不阻塞。最多8种类型同时生成。'),
    ('质量控制',
     '每种资源生成后由QualityAgent检查质量分。分数<0.7且重试<2次时，提升难度重新生成。'
     '最多重试2次。超过重试次数则接受当前结果或降级到WebSearch兜底。'),
    ('WebSearch兜底',
     '当LLM生成失败(超时/异常)时，自动调用web_search_service搜索网络资源。'
     '搜索策略: 视频→Bilibili，流程图→draw.io资源，文档→CSDN/知乎/掘金。'
     '返回格式化的网络资源列表，前端展示为可点击链接。'),
]
for title, desc in act_items:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}: ')
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.add_run(desc)

doc.add_page_break()

# ── 第五章：多智能体协作泳道图 ────────────────────────────────────────
doc.add_heading('五、多智能体协作泳道图', level=1)
doc.add_paragraph(
    '本图展示了LangGraph工作流中多个智能体的协作关系，按职责划分为5个泳道。'
    '展示了资源生成场景下的完整消息流转。'
)
doc.add_picture(os.path.join(OUT, '4_多智能体协作泳道图.png'), width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('泳道职责说明', level=2)
lane_items = [
    ('用户/前端泳道',
     '发送学习请求(POST /learning/chat-stream)，接收SSE流式结果。'
     '前端根据返回的resource_type分发到对应渲染组件。'),
    ('MasterAgent泳道',
     '两阶段意图识别: Stage1关键词匹配(快速,无LLM) → Stage2 LLM分类(备用)。'
     '6种意图: tutoring/resource_generation/learning_path/assessment/exercise/general_chat。'
     '识别后通过route_after_master路由到对应Agent。'),
    ('ProfileAgent/PathAgent泳道',
     'ProfileAgent: 从对话中提取画像信息(extract_profile)。'
     'PathAgent: 规划学习路径(build_learning_path)，使用LLM+图算法修正。'
     'route_after_path决定下一步: learning_path意图→aggregate，resource_generation→resource_planner。'),
    ('ResourcePlanner/生成Agent泳道',
     'ResourcePlanner确定要生成的8种类型。'
     '然后并行启动对应的生成Agent(gen_document, gen_quiz, gen_mindmap等)。'
     '每个生成Agent独立调用agent_runtime.build_learning_resource()。'),
    ('QualityAgent/Recommendation泳道',
     'QualityAgent检查生成质量(分数<0.7重试)。'
     'RecommendationAgent生成推荐列表(含suggested_generation虚拟推荐)。'
     '最终通过aggregate汇总所有结果返回给用户。'),
]
for title, desc in lane_items:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}: ')
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.add_run(desc)

doc.add_page_break()

# ── 第六章：学习路径节点状态图 ────────────────────────────────────────
doc.add_heading('六、学习路径节点状态图', level=1)
doc.add_paragraph(
    '本图展示了学习路径中每个节点的状态转换规则。'
    '节点有5种状态: LOCKED(锁定) → AVAILABLE(可用) → LEARNING(学习中) → COMPLETED(完成) / SKIPPED(跳过)。'
)
doc.add_picture(os.path.join(OUT, '5_学习路径节点状态图.png'), width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('状态转换规则', level=2)
state_items = [
    ('LOCKED → AVAILABLE',
     '前置依赖节点全部完成(COMPLETED)时自动解锁。'
     '系统在complete_path_node()中检查依赖关系，自动将下游LOCKED节点转为AVAILABLE。'),
    ('AVAILABLE → LEARNING',
     '用户点击"开始学习"按钮时触发。前端调用POST /learning-paths/start-node接口，'
     '后端将节点状态从AVAILABLE转为LEARNING。'),
    ('LEARNING → COMPLETED',
     '学习完成后通过测验(accuracy >= 阈值)自动标记完成。'
     '或用户手动标记完成(handleMarkComplete)。完成后自动解锁下游节点。'),
    ('AVAILABLE → SKIPPED',
     '当画像显示该知识点已掌握(composite_score >= 0.7)时自动跳过。'
     '跳过的节点也视为"完成"，可以解锁下游依赖。'),
    ('自适应策略',
     '每个节点有推荐资源类型(recommended_resource_types)，来自策略引擎。'
     'composite_score < 0.4: 补弱(document+mindmap+video)；'
     '0.4-0.7: 巩固(document+quiz+code)；>= 0.7: 进阶(quiz+code+reading)。'),
]
for title, desc in state_items:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}: ')
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.add_run(desc)

# ── 保存 ─────────────────────────────────────────────────────────────
out_path = os.path.join(OUT, 'AutoLearning项目图解说明.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
